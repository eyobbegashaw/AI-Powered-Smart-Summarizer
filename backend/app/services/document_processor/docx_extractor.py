import io
from docx import Document as DocxDocument
from typing import Dict
from loguru import logger


class DocxExtractor:
    def __init__(self):
        pass
    
    async def extract(self, file_bytes: bytes) -> Dict:
        """Extract text and metadata from DOCX file"""
        
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            text = '\n'.join(paragraphs)
            
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            if tables_text:
                text += '\n\n' + '\n'.join(tables_text)
            
            metadata = {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'word_count': len(text.split()),
                'paragraph_count': len(paragraphs)
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {
                'text': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }